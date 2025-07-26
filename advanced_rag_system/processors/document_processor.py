"""
Document Processor - Handles multi-modal document processing and intelligent chunking
"""

import asyncio
import mimetypes
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from dataclasses import dataclass
from loguru import logger

from ..utils.config import ProcessingConfig

@dataclass
class ProcessedDocument:
    """Processed document structure"""
    content: str
    metadata: Dict[str, Any]
    structure: Dict[str, Any]
    chunks: List[Dict[str, Any]]
    processing_info: Dict[str, Any]

class DocumentProcessor:
    """
    Advanced Document Processor with multi-modal support
    
    Features:
    - Multiple format support (PDF, DOCX, TXT, HTML, MD, JSON)
    - Intelligent text extraction
    - Structure preservation
    - OCR capabilities for images
    - Metadata extraction
    - Smart chunking strategies
    """
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.is_initialized = False
        
        # Format handlers
        self.format_handlers = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.html': self._process_html,
            '.htm': self._process_html,
            '.json': self._process_json,
            '.xml': self._process_xml
        }
        
        logger.info("Document Processor initialized")
    
    async def initialize(self):
        """Initialize the document processor"""
        if self.is_initialized:
            return
        
        try:
            # Initialize any required libraries or models
            self.is_initialized = True
            logger.info("Document Processor initialization completed")
            
        except Exception as e:
            logger.error(f"Failed to initialize Document Processor: {str(e)}")
            raise
    
    async def process_document(
        self,
        document_path: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Process a document and extract content, metadata, and structure"""
        if not self.is_initialized:
            await self.initialize()
        
        try:
            document_path = Path(document_path)
            
            # Validate file exists and is supported
            if not document_path.exists():
                raise FileNotFoundError(f"Document not found: {document_path}")
            
            file_extension = document_path.suffix.lower()
            if file_extension not in self.format_handlers:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Extract basic file metadata
            file_metadata = await self._extract_file_metadata(document_path)
            
            # Combine with provided metadata
            combined_metadata = {**file_metadata}
            if metadata:
                combined_metadata.update(metadata)
            
            # Process document based on format
            handler = self.format_handlers[file_extension]
            processed_content = await handler(document_path)
            
            # Extract document structure
            structure = await self._extract_structure(processed_content)
            
            # Create processed document
            result = {
                "content": processed_content,
                "metadata": combined_metadata,
                "structure": structure,
                "processing_info": {
                    "processor_version": "1.0.0",
                    "format": file_extension,
                    "content_length": len(processed_content),
                    "word_count": len(processed_content.split()) if processed_content else 0
                }
            }
            
            logger.info(f"Successfully processed document: {document_path.name}")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {document_path}: {str(e)}")
            raise
    
    async def chunk_text(
        self,
        text: str,
        chunk_size: int = None,
        overlap: int = None,
        preserve_structure: bool = True
    ) -> List[str]:
        """Intelligent text chunking with structure preservation"""
        chunk_size = chunk_size or self.config.chunk_size
        overlap = overlap or self.config.chunk_overlap
        
        try:
            if not text:
                return []
            
            # If text is shorter than chunk size, return as single chunk
            if len(text) <= chunk_size:
                return [text]
            
            chunks = []
            
            if preserve_structure:
                # Try to split on natural boundaries
                chunks = await self._smart_chunking(text, chunk_size, overlap)
            else:
                # Simple sliding window chunking
                chunks = await self._sliding_window_chunking(text, chunk_size, overlap)
            
            logger.info(f"Created {len(chunks)} chunks from text of length {len(text)}")
            return chunks
            
        except Exception as e:
            logger.error(f"Error chunking text: {str(e)}")
            # Fallback to simple chunking
            return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size-overlap)]
    
    async def _extract_file_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract basic file metadata"""
        try:
            stat = file_path.stat()
            mime_type, _ = mimetypes.guess_type(str(file_path))
            
            return {
                "filename": file_path.name,
                "file_path": str(file_path),
                "file_size": stat.st_size,
                "created_time": stat.st_ctime,
                "modified_time": stat.st_mtime,
                "mime_type": mime_type,
                "extension": file_path.suffix.lower()
            }
            
        except Exception as e:
            logger.error(f"Error extracting file metadata: {str(e)}")
            return {"filename": file_path.name}
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Process PDF documents"""
        try:
            import pypdf
            
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
            
            content = "\n\n".join(text_content)
            
            # If no text extracted and OCR is enabled, try OCR
            if not content.strip() and self.config.enable_ocr:
                content = await self._ocr_pdf(file_path)
            
            return content
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            if self.config.enable_ocr:
                return await self._ocr_pdf(file_path)
            return ""
    
    async def _process_docx(self, file_path: Path) -> str:
        """Process DOCX documents"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("\n".join(table_text))
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            return ""
    
    async def _process_text(self, file_path: Path) -> str:
        """Process plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors
            with open(file_path, 'rb') as file:
                return file.read().decode('utf-8', errors='replace')
                
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            return ""
    
    async def _process_markdown(self, file_path: Path) -> str:
        """Process Markdown files"""
        try:
            # First get the raw text
            raw_content = await self._process_text(file_path)
            
            # Optionally convert markdown to plain text while preserving structure
            # For now, return raw content
            return raw_content
            
        except Exception as e:
            logger.error(f"Error processing Markdown: {str(e)}")
            return ""
    
    async def _process_html(self, file_path: Path) -> str:
        """Process HTML files"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            logger.error(f"Error processing HTML: {str(e)}")
            return ""
    
    async def _process_json(self, file_path: Path) -> str:
        """Process JSON files"""
        try:
            import json
            
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert JSON to readable text format
            if isinstance(data, dict):
                text_parts = []
                for key, value in data.items():
                    if isinstance(value, (str, int, float)):
                        text_parts.append(f"{key}: {value}")
                    elif isinstance(value, list):
                        text_parts.append(f"{key}: {', '.join(map(str, value))}")
                    else:
                        text_parts.append(f"{key}: {json.dumps(value, indent=2)}")
                
                return "\n".join(text_parts)
            else:
                return json.dumps(data, indent=2)
                
        except Exception as e:
            logger.error(f"Error processing JSON: {str(e)}")
            return ""
    
    async def _process_xml(self, file_path: Path) -> str:
        """Process XML files"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'xml')
            
            # Extract text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)
            
            return text
            
        except Exception as e:
            logger.error(f"Error processing XML: {str(e)}")
            return ""
    
    async def _ocr_pdf(self, file_path: Path) -> str:
        """Perform OCR on PDF (placeholder implementation)"""
        try:
            # This would require additional OCR libraries like Tesseract
            logger.info(f"OCR processing would be performed on {file_path}")
            return "OCR processing not implemented in this version"
            
        except Exception as e:
            logger.error(f"Error in OCR processing: {str(e)}")
            return ""
    
    async def _extract_structure(self, content: str) -> Dict[str, Any]:
        """Extract document structure information"""
        try:
            structure = {
                "headings": [],
                "paragraphs": 0,
                "lists": 0,
                "tables": 0,
                "sections": []
            }
            
            lines = content.split('\n')
            
            # Simple structure detection
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Detect headings (lines that are all caps or start with #)
                if line.isupper() or line.startswith('#'):
                    structure["headings"].append(line)
                
                # Count paragraphs (non-empty lines)
                if line and not line.startswith(('#', '-', '*', '1.', '2.')):
                    structure["paragraphs"] += 1
                
                # Detect lists
                if line.startswith(('-', '*', '•')) or (line[0].isdigit() and '.' in line[:5]):
                    structure["lists"] += 1
            
            return structure
            
        except Exception as e:
            logger.error(f"Error extracting structure: {str(e)}")
            return {"headings": [], "paragraphs": 0, "lists": 0, "tables": 0, "sections": []}
    
    async def _smart_chunking(
        self,
        text: str,
        chunk_size: int,
        overlap: int
    ) -> List[str]:
        """Smart chunking that respects natural boundaries"""
        try:
            chunks = []
            
            # Split on paragraphs first
            paragraphs = text.split('\n\n')
            
            current_chunk = ""
            
            for paragraph in paragraphs:
                # If adding this paragraph would exceed chunk size
                if len(current_chunk) + len(paragraph) > chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                        
                        # Start new chunk with overlap
                        if overlap > 0 and len(current_chunk) > overlap:
                            current_chunk = current_chunk[-overlap:] + "\n\n" + paragraph
                        else:
                            current_chunk = paragraph
                    else:
                        # Paragraph itself is too long, split it
                        sentences = self._split_into_sentences(paragraph)
                        for sentence in sentences:
                            if len(current_chunk) + len(sentence) > chunk_size:
                                if current_chunk:
                                    chunks.append(current_chunk.strip())
                                    current_chunk = sentence
                                else:
                                    # Sentence itself is too long, force split
                                    chunks.append(sentence[:chunk_size])
                                    current_chunk = sentence[chunk_size-overlap:]
                            else:
                                current_chunk += " " + sentence if current_chunk else sentence
                else:
                    current_chunk += "\n\n" + paragraph if current_chunk else paragraph
            
            # Add the last chunk
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error in smart chunking: {str(e)}")
            return await self._sliding_window_chunking(text, chunk_size, overlap)
    
    async def _sliding_window_chunking(
        self,
        text: str,
        chunk_size: int,
        overlap: int
    ) -> List[str]:
        """Simple sliding window chunking"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            
            if chunk.strip():
                chunks.append(chunk)
            
            start = end - overlap
            
            # Prevent infinite loop
            if start >= end:
                break
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Simple sentence splitting"""
        import re
        
        # Simple sentence splitting on periods, exclamation marks, and question marks
        sentences = re.split(r'[.!?]+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    async def get_processing_stats(self) -> Dict[str, Any]:
        """Get document processing statistics"""
        return {
            "supported_formats": list(self.format_handlers.keys()),
            "config": {
                "chunk_size": self.config.chunk_size,
                "chunk_overlap": self.config.chunk_overlap,
                "enable_ocr": self.config.enable_ocr,
                "enable_multimodal": self.config.enable_multimodal
            },
            "initialized": self.is_initialized
        }