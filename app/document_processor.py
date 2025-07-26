import os
import hashlib
import json
from typing import List, Dict, Any, Optional
from pathlib import Path
import aiofiles
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import pandas as pd
from bs4 import BeautifulSoup
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
import tiktoken
from app.config import settings


class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        self.encoding = tiktoken.get_encoding("cl100k_base")
    
    async def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extract text content from various file types"""
        try:
            if file_type.lower() == 'pdf':
                return await self._extract_from_pdf(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                return await self._extract_from_docx(file_path)
            elif file_type.lower() in ['xlsx', 'xls', 'csv']:
                return await self._extract_from_excel(file_path)
            elif file_type.lower() in ['txt', 'md', 'py', 'js', 'html', 'css', 'json', 'xml']:
                return await self._extract_from_text(file_path)
            elif file_type.lower() == 'html':
                return await self._extract_from_html(file_path)
            else:
                # Try to read as text file
                return await self._extract_from_text(file_path)
        except Exception as e:
            raise ValueError(f"Error extracting text from {file_type} file: {str(e)}")
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            raise ValueError(f"Error reading PDF: {str(e)}")
        return text.strip()
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {str(e)}")
    
    async def _extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel/CSV files"""
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Convert DataFrame to text representation
            text = f"Columns: {', '.join(df.columns.tolist())}\n\n"
            text += df.to_string(index=False)
            return text
        except Exception as e:
            raise ValueError(f"Error reading Excel/CSV: {str(e)}")
    
    async def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = await file.read()
            return content
        except Exception as e:
            raise ValueError(f"Error reading text file: {str(e)}")
    
    async def _extract_from_html(self, file_path: str) -> str:
        """Extract text from HTML files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = await file.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text()
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            return text
        except Exception as e:
            raise ValueError(f"Error reading HTML: {str(e)}")
    
    def create_chunks(self, text: str, metadata: Dict[str, Any] = None) -> List[LangchainDocument]:
        """Split text into chunks for vector storage"""
        if not text.strip():
            return []
        
        # Create documents with metadata
        docs = [LangchainDocument(page_content=text, metadata=metadata or {})]
        
        # Split into chunks
        chunks = self.text_splitter.split_documents(docs)
        
        # Add chunk-specific metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata.update({
                'chunk_index': i,
                'chunk_size': len(chunk.page_content),
                'token_count': len(self.encoding.encode(chunk.page_content))
            })
        
        return chunks
    
    def calculate_content_hash(self, content: str) -> str:
        """Calculate SHA-256 hash of content"""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()
    
    def extract_metadata(self, file_path: str, content: str) -> Dict[str, Any]:
        """Extract metadata from file and content"""
        file_stat = os.stat(file_path)
        
        metadata = {
            'file_size': file_stat.st_size,
            'created_time': file_stat.st_ctime,
            'modified_time': file_stat.st_mtime,
            'content_length': len(content),
            'token_count': len(self.encoding.encode(content)),
            'line_count': content.count('\n') + 1,
            'word_count': len(content.split()),
            'character_count': len(content)
        }
        
        return metadata
    
    def create_content_preview(self, content: str, max_length: int = 500) -> str:
        """Create a preview of the content"""
        if len(content) <= max_length:
            return content
        
        # Try to cut at sentence boundary
        preview = content[:max_length]
        last_sentence = preview.rfind('.')
        if last_sentence > max_length * 0.7:  # If we have a reasonable sentence boundary
            preview = preview[:last_sentence + 1]
        
        return preview + "..."
    
    async def analyze_content_type(self, content: str) -> Dict[str, Any]:
        """Analyze content to determine its type and characteristics"""
        analysis = {
            'content_type': 'text',
            'language': 'unknown',
            'structure': 'unstructured',
            'topics': [],
            'entities': [],
            'sentiment': 'neutral'
        }
        
        # Basic content type detection
        if any(tag in content.lower() for tag in ['<html>', '<div>', '<p>', '<table>']):
            analysis['content_type'] = 'html'
            analysis['structure'] = 'markup'
        elif content.strip().startswith('{') and content.strip().endswith('}'):
            analysis['content_type'] = 'json'
            analysis['structure'] = 'structured'
        elif ',' in content and '\n' in content:
            lines = content.split('\n')
            if len(lines) > 1 and len(lines[0].split(',')) > 1:
                analysis['content_type'] = 'csv'
                analysis['structure'] = 'tabular'
        
        # Simple keyword extraction (in a real implementation, you might use NLP libraries)
        words = content.lower().split()
        word_freq = {}
        for word in words:
            if len(word) > 3:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Get top keywords as topics
        analysis['topics'] = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
        
        return analysis


# Global processor instance
document_processor = DocumentProcessor()